"""
ResumeRenderer — deterministic DOCX generation via python-docx.

Architecture decision: the LLM produces structured JSON; this module turns it
into a stable Word document. No formatting logic lives in the LLM prompt —
all layout decisions (borders, spacing, fonts, table structure) live here.

Customization: see config.py for every knob.
"""

from __future__ import annotations

from pathlib import Path
from typing import List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

# ── local imports ────────────────────────────────────────────────────────────
from .config import RenderConfig
from ..schemas.resume_schema import (
    EducationEntry,
    ExperienceEntry,
    ProjectEntry,
    SkillsData,
    TailoredResume,
)


# ─────────────────────────────────────────────────────────────────────────────
# Low-level XML helpers
# ─────────────────────────────────────────────────────────────────────────────

def _hex_to_rgb(hex_color: str):
    """Convert 6-char hex string (no #) to (r, g, b) tuple."""
    h = hex_color.lstrip("#")
    return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))


def _set_run_color(run, hex_color: str):
    r, g, b = _hex_to_rgb(hex_color)
    run.font.color.rgb = RGBColor(r, g, b)


def _set_paragraph_border_bottom(paragraph, color: str = "1F3864", size: int = 12):
    """
    Draw a bottom border under a paragraph (used for section title dividers).

    CUSTOMIZE:
      color — hex string, controls divider line color
      size  — in half-points; 12 = 1.5pt, 6 = 0.75pt, 4 = 0.5pt
    """
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color.lstrip("#"))
    pBdr.append(bottom)
    pPr.append(pBdr)


def _clear_table_borders(table):
    """
    Remove all visible borders from a table (invisible layout grid).

    CUSTOMIZE: To show inner grid lines instead, call _set_cell_borders on
    individual cells rather than using this function.
    """
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)

    tblBorders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        tblBorders.append(el)

    # Replace any existing tblBorders
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(tblBorders)


def _set_cell_borders(cell, top=False, bottom=False, left=False, right=False,
                      color="CCCCCC", size=4):
    """
    Set selective visible borders on a table cell.

    CUSTOMIZE:
      Set top/bottom/left/right to True to show that edge.
      color — hex border color
      size  — half-points
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    existing = tcPr.find(qn("w:tcBorders"))
    if existing is not None:
        tcPr.remove(existing)

    tcBorders = OxmlElement("w:tcBorders")
    edges = {"top": top, "bottom": bottom, "start": left, "end": right}
    for edge, show in edges.items():
        el = OxmlElement(f"w:{edge}")
        if show:
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), str(size))
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), color.lstrip("#"))
        else:
            el.set(qn("w:val"), "none")
            el.set(qn("w:sz"), "0")
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), "auto")
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _no_cell_spacing(cell):
    """Remove default cell margins (for tighter layout control)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side in ("top", "bottom", "start", "end"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), "0")
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tcPr.append(tcMar)


def _add_hyperlink(paragraph, text: str, url: str, font_name: str,
                   font_size: int, color_hex: str = "0563C1"):
    """
    Insert a clickable hyperlink run into an existing paragraph.

    CUSTOMIZE:
      color_hex — link color; default is standard Word blue
      Set underline inside the run properties to False to remove underline.
    """
    from docx.opc.constants import RELATIONSHIP_TYPE as RT  # noqa

    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    # Font
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rPr.append(rFonts)

    # Size (half-points)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(font_size * 2))
    szCs = OxmlElement("w:szCs")
    szCs.set(qn("w:val"), str(font_size * 2))
    rPr.append(sz)
    rPr.append(szCs)

    # Color
    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), color_hex.lstrip("#"))
    rPr.append(color_el)

    # Underline
    u_el = OxmlElement("w:u")
    u_el.set(qn("w:val"), "single")
    rPr.append(u_el)

    run.append(rPr)

    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    run.append(t)

    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink


def _add_plain_run(paragraph, text: str, font_name: str, font_size: int,
                   bold: bool = False, color_hex: str = "000000",
                   preserve_space: bool = True):
    """Add a plain styled text run to a paragraph."""
    run = paragraph.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    _set_run_color(run, color_hex)
    if preserve_space:
        run._r.get_or_add_rPr()  # ensure rPr exists
    return run


def _para_spacing(paragraph, before_pt: int = 0, after_pt: int = 0):
    """Set paragraph spacing (before/after in pt). Line spacing inherits from style."""
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before_pt)
    pf.space_after = Pt(after_pt)


# ─────────────────────────────────────────────────────────────────────────────
# ResumeRenderer
# ─────────────────────────────────────────────────────────────────────────────

class ResumeRenderer:
    """
    Converts a TailoredResume (or equivalent dict) into a .docx file.

    Usage:
        renderer = ResumeRenderer()                        # default config
        renderer = ResumeRenderer(config=my_config)        # custom config
        renderer.render(resume, "output/my_resume.docx")
    """

    def __init__(self, config: Optional[RenderConfig] = None):
        self.cfg = config or RenderConfig()

    # ── Public API ────────────────────────────────────────────────────────

    def render(self, resume: TailoredResume, output_path: str) -> str:
        """Render and save the resume. Returns the output path."""
        self.doc = Document()
        self._setup_document()

        # Determine section order (resume JSON overrides config default)
        order = resume.section_order or self.cfg.default_section_order

        for section in order:
            if section == "header":
                self._render_header(resume.header)
            elif section == "summary" and resume.summary:
                self._render_summary(resume.summary)
            elif section == "skills" and self._has_skills(resume.skills):
                self._render_skills(resume.skills)
            elif section == "experience" and resume.experience:
                self._render_experience(resume.experience)
            elif section == "projects" and resume.projects:
                self._render_projects(resume.projects)
            elif section == "education" and resume.education:
                self._render_education(resume.education)
            elif section == "certifications" and resume.certifications:
                self._render_certifications(resume.certifications)
            elif section == "awards" and resume.awards:
                self._render_awards(resume.awards)

        out = Path(output_path)
        out.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(out))
        return str(out)

    # ── Document setup ────────────────────────────────────────────────────

    def _setup_document(self):
        cfg = self.cfg
        doc = self.doc

        # Page size and margins
        for section in doc.sections:
            section.page_width = Inches(cfg.page_width_inches)
            section.page_height = Inches(cfg.page_height_inches)
            section.top_margin = Inches(cfg.margin_top_inches)
            section.bottom_margin = Inches(cfg.margin_bottom_inches)
            section.left_margin = Inches(cfg.margin_left_inches)
            section.right_margin = Inches(cfg.margin_right_inches)

        # Default paragraph style
        style = doc.styles["Normal"]
        style.font.name = cfg.font_name
        style.font.size = Pt(cfg.body_font_size)

        # python-docx inserts one empty paragraph when Document() is created.
        # Remove it so the name heading starts exactly at the top margin
        # instead of being pushed down by an invisible blank line.
        if doc.paragraphs and doc.paragraphs[0].text == "":
            p = doc.paragraphs[0]._element
            p.getparent().remove(p)

    # ── Section: Header ───────────────────────────────────────────────────

    def _render_header(self, header):
        cfg = self.cfg

        # ── Name ─────────────────────────────────────────────────────────
        name_para = self.doc.add_paragraph()
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _para_spacing(name_para, before_pt=0, after_pt=2)
        name_run = name_para.add_run(header.name)
        name_run.font.name = cfg.font_name_heading
        name_run.font.size = Pt(cfg.name_font_size)
        name_run.font.bold = True
        _set_run_color(name_run, cfg.name_color)

        # ── Contact line ─────────────────────────────────────────────────
        contact_para = self.doc.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _para_spacing(contact_para, before_pt=0, after_pt=4)

        parts = []
        if header.phone:
            parts.append(("plain", header.phone, ""))
        if header.email:
            parts.append(("link", header.email, f"mailto:{header.email}"))
        if header.location:
            parts.append(("plain", header.location, ""))
        if header.linkedin:
            label = "LinkedIn"
            parts.append(("link", label, header.linkedin))
        if header.github:
            label = "GitHub"
            parts.append(("link", label, header.github))
        if header.website:
            parts.append(("link", header.website, header.website))

        for i, (kind, text, url) in enumerate(parts):
            if i > 0:
                sep = _add_plain_run(
                    contact_para, "  |  ",
                    cfg.font_name, cfg.contact_font_size,
                    color_hex=cfg.body_color,
                )
            if kind == "link" and url:
                _add_hyperlink(
                    contact_para, text, url,
                    cfg.font_name, cfg.contact_font_size,
                    cfg.link_color,
                )
            else:
                _add_plain_run(
                    contact_para, text,
                    cfg.font_name, cfg.contact_font_size,
                    color_hex=cfg.body_color,
                )

    # ── Section: Summary ─────────────────────────────────────────────────

    def _render_summary(self, summary: str):
        self._render_section_title("Summary")
        p = self.doc.add_paragraph()
        _para_spacing(p, before_pt=self.cfg.bullet_spacing_before,
                      after_pt=self.cfg.bullet_spacing_after)
        _add_plain_run(p, summary, self.cfg.font_name, self.cfg.body_font_size,
                       color_hex=self.cfg.body_color)

    # ── Section: Skills ──────────────────────────────────────────────────

    def _render_skills(self, skills: SkillsData):
        cfg = self.cfg
        self._render_section_title("Skills")

        rows = []
        if skills.core:
            rows.append(("Core", " · ".join(skills.core)))
        if skills.tools:
            rows.append(("Tools", " · ".join(skills.tools)))
        if skills.methods:
            rows.append(("Methods", " · ".join(skills.methods)))
        if skills.domains:
            rows.append(("Domains", " · ".join(skills.domains)))

        for label, values in rows:
            p = self.doc.add_paragraph()
            _para_spacing(p, before_pt=0, after_pt=2)
            label_run = p.add_run(f"{label}: ")
            label_run.font.name = cfg.font_name
            label_run.font.size = Pt(cfg.body_font_size)
            label_run.font.bold = True
            _set_run_color(label_run, cfg.body_color)
            value_run = p.add_run(values)
            value_run.font.name = cfg.font_name
            value_run.font.size = Pt(cfg.body_font_size)
            value_run.font.bold = False
            _set_run_color(value_run, cfg.body_color)

    # ── Section: Experience ───────────────────────────────────────────────

    def _render_experience(self, experience: List[ExperienceEntry]):
        self._render_section_title("Experience")
        for i, job in enumerate(experience):
            self._render_role_heading(
                left_top=job.title,
                left_bottom=job.company + (f", {job.location}" if job.location else ""),
                right_top=job.dates,
                right_bottom="",
                spacing_before=self.cfg.role_spacing_before if i > 0 else 0,
            )
            for bullet in job.bullets:
                self._render_bullet(bullet)
            # optional project-style links on experience entries
            for lk in job.links:
                if lk.url:
                    p = self.doc.add_paragraph()
                    _para_spacing(p, before_pt=0, after_pt=2)
                    p.paragraph_format.left_indent = Inches(self.cfg.bullet_indent_left)
                    _add_hyperlink(p, lk.label or lk.url, lk.url,
                                   self.cfg.font_name, self.cfg.body_font_size,
                                   self.cfg.link_color)

    # ── Section: Projects ─────────────────────────────────────────────────

    def _render_projects(self, projects: List[ProjectEntry]):
        self._render_section_title("Projects")
        for i, proj in enumerate(projects):
            subtitle = proj.subtitle or ""
            left_bottom = subtitle
            self._render_role_heading(
                left_top=proj.name,
                left_bottom=left_bottom,
                right_top=proj.dates,
                right_bottom=proj.location,
                spacing_before=self.cfg.role_spacing_before if i > 0 else 0,
                left_top_url=proj.url or "",
            )
            for bullet in proj.bullets:
                self._render_bullet(bullet)

    # ── Section: Education ────────────────────────────────────────────────

    def _render_education(self, education: List[EducationEntry]):
        self._render_section_title("Education")
        for i, edu in enumerate(education):
            self._render_role_heading(
                left_top=edu.school,
                left_bottom=edu.degree,
                right_top=edu.dates,
                right_bottom=edu.location,
                spacing_before=self.cfg.role_spacing_before if i > 0 else 0,
            )
            for detail in edu.details:
                self._render_bullet(detail)

    # ── Section: Certifications ───────────────────────────────────────────

    def _render_certifications(self, certs: List[str]):
        self._render_section_title("Certifications")
        for cert in certs:
            self._render_bullet(cert)

    # ── Section: Awards ───────────────────────────────────────────────────

    def _render_awards(self, awards: List[str]):
        self._render_section_title("Awards")
        for award in awards:
            self._render_bullet(award)

    # ── Layout primitives ─────────────────────────────────────────────────

    def _render_section_title(self, title: str):
        """
        Section heading paragraph with optional bottom border line.

        CUSTOMIZE via config:
          show_section_bottom_border  → toggle the divider line
          section_border_color        → line color
          section_border_size         → line weight (half-points)
          section_title_color         → text color
          section_title_size          → font size
          section_spacing_before/after → vertical breathing room
        """
        cfg = self.cfg
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _para_spacing(p,
                      before_pt=cfg.section_spacing_before,
                      after_pt=cfg.section_spacing_after)

        run = p.add_run(title.upper())
        run.font.name = cfg.font_name_heading
        run.font.size = Pt(cfg.section_title_size)
        run.font.bold = True
        _set_run_color(run, cfg.section_title_color)

        if cfg.show_section_bottom_border:
            _set_paragraph_border_bottom(
                p,
                color=cfg.section_border_color,
                size=cfg.section_border_size,
            )

    def _render_role_heading(
        self,
        left_top: str,
        left_bottom: str,
        right_top: str,
        right_bottom: str,
        spacing_before: int = 0,
        left_top_url: str = "",
    ):
        """
        Two-column heading row using a borderless 2-cell table.

        Layout:
          [left_top (bold)]           [right_top (right-aligned)]
          [left_bottom (italic)]      [right_bottom (right-aligned)]

        CUSTOMIZE:
          Table borders are always invisible (layout only).
          To add a visible separator above each role, set a top cell border
          in _set_cell_borders() calls below.
        """
        cfg = self.cfg

        # Content width = page width minus margins
        content_w = cfg.page_width_inches - cfg.margin_left_inches - cfg.margin_right_inches
        left_w = content_w * 0.68
        right_w = content_w * 0.32

        table = self.doc.add_table(rows=1, cols=2)
        _clear_table_borders(table)
        table.autofit = False

        left_cell = table.cell(0, 0)
        right_cell = table.cell(0, 1)

        left_cell.width = Inches(left_w)
        right_cell.width = Inches(right_w)

        # ── Left cell ────────────────────────────────────────────────────
        lp = left_cell.paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _para_spacing(lp, before_pt=spacing_before, after_pt=0)

        if left_top_url:
            _add_hyperlink(lp, left_top, left_top_url,
                           cfg.font_name, cfg.role_title_size, cfg.link_color)
        else:
            r = lp.add_run(left_top)
            r.font.name = cfg.font_name
            r.font.size = Pt(cfg.role_title_size)
            r.font.bold = True
            _set_run_color(r, cfg.body_color)

        if left_bottom:
            lp2 = left_cell.add_paragraph()
            lp2.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _para_spacing(lp2, before_pt=0, after_pt=0)
            r2 = lp2.add_run(left_bottom)
            r2.font.name = cfg.font_name
            r2.font.size = Pt(cfg.body_font_size)
            r2.font.bold = False
            r2.font.italic = True
            _set_run_color(r2, cfg.body_color)

        # ── Right cell ───────────────────────────────────────────────────
        rp = right_cell.paragraphs[0]
        rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _para_spacing(rp, before_pt=spacing_before, after_pt=0)

        if right_top:
            r = rp.add_run(right_top)
            r.font.name = cfg.font_name
            r.font.size = Pt(cfg.body_font_size)
            r.font.bold = False
            _set_run_color(r, cfg.body_color)

        if right_bottom:
            rp2 = right_cell.add_paragraph()
            rp2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            _para_spacing(rp2, before_pt=0, after_pt=0)
            r2 = rp2.add_run(right_bottom)
            r2.font.name = cfg.font_name
            r2.font.size = Pt(cfg.body_font_size)
            _set_run_color(r2, cfg.body_color)

        # Remove cell padding
        _no_cell_spacing(left_cell)
        _no_cell_spacing(right_cell)

    def _render_bullet(self, text: str):
        """
        Add a bullet paragraph.

        CUSTOMIZE via config:
          bullet_indent_left  → left indent (inches)
          bullet_hanging      → hanging indent for the bullet character
          bullet_spacing_before/after → vertical spacing
          body_font_size      → bullet text size
        """
        cfg = self.cfg
        p = self.doc.add_paragraph(style="List Bullet")
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        pf = p.paragraph_format
        pf.left_indent = Inches(cfg.bullet_indent_left)
        pf.first_line_indent = Inches(-cfg.bullet_hanging)
        pf.space_before = Pt(cfg.bullet_spacing_before)
        pf.space_after = Pt(cfg.bullet_spacing_after)

        run = p.add_run(text)
        run.font.name = cfg.font_name
        run.font.size = Pt(cfg.body_font_size)
        _set_run_color(run, cfg.body_color)
        return p

    # ── Helpers ───────────────────────────────────────────────────────────

    def _has_skills(self, skills: SkillsData) -> bool:
        return any([skills.core, skills.tools, skills.methods, skills.domains])
